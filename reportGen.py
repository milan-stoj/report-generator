import os
from docx import Document

def reportGen(self):
    tests = []
    for index in range(self.listWidget.count()):
        tests.append(self.listWidget.item(index).text())
    
    document = Document()
    
    document.add_heading(self.entryTitle.text(), 0)
    document.add_heading(("Prepared By: ", self.entryAuthor.text()),1)
    paragraph = document.add_paragraph('Normal text, ')
    paragraph.add_run('text with emphasis.', 'Emphasis')
    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    # get table data -------------

    table = document.add_table(1, 1)
    table.style = 'Colorful List'
    # populate header row --------
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Recommended Testing Procedures'
    
    # add a data row for each item
    for item in tests:
        cells = table.add_row().cells
        cells[0].text = item
   
    document.save('test.docx')
    os.startfile('test.docx')
        
